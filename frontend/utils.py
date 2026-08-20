import easyocr
import numpy as np
from PIL import Image
import io
import re
import datetime
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

reader = None

def extract_ocr_data(image_bytes):
    global reader
    if reader is None:
        reader = easyocr.Reader(['en'])
    img = Image.open(io.BytesIO(image_bytes)).convert('RGB')
    img_np = np.array(img)
    results = reader.readtext(img_np, detail=0)
    text = " ".join(results)
    
    extracted = {}
    name_match = re.search(r'Name[:\s]+([A-Za-z\s]+)(?:Age|Gender|$)', text, re.IGNORECASE)
    if name_match: extracted['name'] = name_match.group(1).strip()
    
    income_match = re.search(r'(?:Income|Salary)[:\s]+(?:Rs\.?|INR|₹)?\s*([\d,]+)', text, re.IGNORECASE)
    if income_match: extracted['gross_monthly_income'] = int(income_match.group(1).replace(',', ''))
    
    loan_match = re.search(r'(?:Loan|Amount)[:\s]+(?:Rs\.?|INR|₹)?\s*([\d,]+)', text, re.IGNORECASE)
    if loan_match: extracted['loan_amount'] = int(loan_match.group(1).replace(',', ''))
    
    return extracted


# -------------------------------------------------------------
# Enterprise Publication-Grade Word (.docx) Generator
# -------------------------------------------------------------

def set_cell_background(cell, fill_hex: str):
    """Sets background shading of a Word table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    for shd in tcPr.xpath('w:shd'):
        tcPr.remove(shd)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    """Sets inner cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    for tcMar in tcPr.xpath('w:tcMar'):
        tcPr.remove(tcMar)
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, top="003366", bottom="003366", insideH="E2E8F0"):
    """Applies clean horizontal table borders with no vertical clutter."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="6" w:space="0" w:color="{top}"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="{bottom}"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{insideH}"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def make_row_cant_split(row):
    """Prevents a table row from splitting across page breaks."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def make_row_header(row):
    """Marks a table row as a repeating header on new pages."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def clean_section_title(raw_text: str) -> str:
    """Removes informal web emojis and standardizes section headers."""
    text = raw_text
    # Remove emoji symbols
    text = re.sub(r'[\U00010000-\U0010ffff]', '', text)
    text = re.sub(r'[\u2600-\u27ff]', '', text)
    text = text.replace('📊', '').replace('👤', '').replace('💰', '').replace('🏢', '').replace('🧠', '').replace('📜', '').replace('📚', '').replace('🏛️', '').replace('📑', '').replace('⚖️', '').replace('📈', '').replace('✍️', '')
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Capitalize section titles formally
    if text.startswith("1.") or "Executive Summary" in text:
        return "1. EXECUTIVE SUMMARY"
    elif text.startswith("2.") or "Applicant Profile" in text:
        return "2. APPLICANT PROFILE & CREDIT FACILITY DETAILS"
    elif text.startswith("3.1") or "MSE Scoring" in text:
        return text.upper() if text.startswith("3.1") else "3.1 CENTRAL BANK OF INDIA MSE SCORING MODEL"
    elif text.startswith("3.") or "Financial Capacity" in text:
        return "3. FINANCIAL CAPACITY & OBLIGATION ASSESSMENT"
    elif text.startswith("4.") or "Predictive Risk" in text:
        return "4. PREDICTIVE RISK & DEFAULT PROBABILITY ASSESSMENT"
    elif text.startswith("5.") or "Policy Adherence" in text:
        return "5. STATUTORY POLICY ADHERENCE & FINAL JUSTIFICATION"
    elif text.startswith("6.") or "References" in text:
        return "6. REGULATORY REFERENCES & POLICY BIBLIOGRAPHY"
    return text

def sanitize_text(text: str) -> str:
    """Completely purges all HTML span, br, and div tags into clean text."""
    clean = re.sub(r'<span[^>]*style="[^"]*color:\s*(?:green|#2e7d32|#4caf50|#15803d)[^"]*"[^>]*>(.*?)</span>', r'\1', text, flags=re.IGNORECASE)
    clean = re.sub(r'<span[^>]*style="[^"]*color:\s*(?:red|#c62828|#d32f2f|#b91c1c)[^"]*"[^>]*>(.*?)</span>', r'\1', clean, flags=re.IGNORECASE)
    clean = re.sub(r'<span[^>]*>(.*?)</span>', r'\1', clean, flags=re.IGNORECASE)
    clean = re.sub(r'<br\s*/?>', '\n', clean, flags=re.IGNORECASE)
    clean = re.sub(r'<[^>]+>', '', clean) # Purge any remaining HTML tags
    # Clean emoji artifacts
    clean = clean.replace('🏛️', '').replace('🏢', '').replace('📊', '').replace('👤', '').replace('💰', '').replace('🧠', '').replace('📜', '').replace('📚', '')
    return clean

def add_formatted_runs(paragraph, text: str, base_font_size=Pt(9.5), default_color=RGBColor(30, 41, 59)):
    """
    Parses inline markdown (**bold**, *italic*, `code`) into Word runs
    and applies executive financial color coding.
    """
    text = sanitize_text(text)
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for token in tokens:
        if not token:
            continue
            
        is_bold = False
        is_italic = False
        is_code = False
        content = token
        
        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            content = token[2:-2]
            is_bold = True
        elif token.startswith('*') and token.endswith('*') and len(token) >= 2:
            content = token[1:-1]
            is_italic = True
        elif token.startswith('`') and token.endswith('`') and len(token) >= 2:
            content = token[1:-1]
            is_code = True
            
        run = paragraph.add_run(content)
        run.bold = is_bold
        run.italic = is_italic
        run.font.size = base_font_size
        
        # Color coding for financial terms and status
        c_upper = content.strip().upper()
        if is_code:
            run.font.name = "Consolas"
            run.font.size = Pt(9.0)
            run.font.color.rgb = RGBColor(194, 65, 12)
        elif c_upper in ["APPROVED", "COMPLIANT", "GRADE A", "CBI 1", "CBI 2", "CBI 3", "CBI 4", "LOW RISK", "VERY LOW RISK", "PRIME", "PROMPT", "TIMELY", "STANDARD ASSET", "HURDLE RATE MET (> 50 MARKS)"]:
            run.bold = True
            run.font.color.rgb = RGBColor(21, 128, 61) # Forest Green
        elif c_upper in ["REJECTED", "NON-COMPLIANT", "GRADE C", "CBI 7", "CBI 8", "CBI 9", "CBI 10", "HIGH RISK", "CRITICAL", "OVERDUE", "BREACH", "VIOLATION", "SUB-HURDLE RATE (<= 50 MARKS)"]:
            run.bold = True
            run.font.color.rgb = RGBColor(185, 28, 28) # Crimson Red
        elif c_upper in ["GRADE B", "CBI 5", "CBI 6", "MODERATE", "MODERATE RISK", "ELEVATED RISK", "DELAYED"]:
            run.bold = True
            run.font.color.rgb = RGBColor(30, 64, 175) # Navy Blue
        else:
            run.font.color.rgb = default_color


def generate_docx(markdown_text: str) -> bytes:
    """
    Generates a publication-grade, beautifully formatted Word (.docx)
    appraisal memorandum with pixel-perfect tables and corporate bank branding.
    """
    doc = Document()
    
    # 1. Page Margins & Setup (0.7" for balanced density)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        
        # Header
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        h_run = header_p.add_run("सेन्ट्रल बैंक ऑफ़ इंडिया  |  CENTRAL BANK OF INDIA  —  CREDIT APPRAISAL MEMORANDUM")
        h_run.font.size = Pt(8.0)
        h_run.font.color.rgb = RGBColor(100, 116, 139)
        
        # Footer
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f_run = footer_p.add_run("Strictly Confidential  •  For Internal Sanction Committee & Audit Use Only")
        f_run.font.size = Pt(8.0)
        f_run.font.italic = True
        f_run.font.color.rgb = RGBColor(148, 163, 184)

    # 2. Base Typography
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    doc.styles['Normal'].font.color.rgb = RGBColor(30, 41, 59)
    
    # 3. Bank Letterhead & Header Card
    banner_table = doc.add_table(rows=1, cols=1)
    banner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner_table.autofit = False
    b_cell = banner_table.cell(0, 0)
    b_cell.width = Inches(7.1)
    set_cell_background(b_cell, "003366") # Corporate Deep Navy
    set_cell_margins(b_cell, top=160, bottom=160, left=200, right=200)
    
    bp = b_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after = Pt(2)
    
    r1 = bp.add_run("सेन्ट्रल बैंक ऑफ़ इंडिया / CENTRAL BANK OF INDIA\n")
    r1.font.size = Pt(14)
    r1.bold = True
    r1.font.color.rgb = RGBColor(255, 255, 255)
    
    r2 = bp.add_run("CREDIT APPRAISAL & SANCTION ASSESSMENT MEMORANDUM")
    r2.font.size = Pt(10.5)
    r2.bold = True
    r2.font.color.rgb = RGBColor(253, 224, 71) # Gold Accent
    
    # 3.1 Metadata Bar below banner
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    set_table_borders(meta_table, top="003366", bottom="CBD5E1", insideH="CBD5E1")
    
    m_cell0 = meta_table.cell(0, 0)
    m_cell0.width = Inches(3.55)
    set_cell_background(m_cell0, "F8FAFC")
    set_cell_margins(m_cell0, top=60, bottom=60, left=100, right=100)
    p_meta0 = m_cell0.paragraphs[0]
    p_meta0.paragraph_format.space_before = Pt(0)
    p_meta0.paragraph_format.space_after = Pt(0)
    r_m0 = p_meta0.add_run(f"DATE OF APPRAISAL: {datetime.date.today().strftime('%d-%b-%Y')}")
    r_m0.font.size = Pt(8.5)
    r_m0.bold = True
    r_m0.font.color.rgb = RGBColor(71, 85, 105)
    
    m_cell1 = meta_table.cell(0, 1)
    m_cell1.width = Inches(3.55)
    set_cell_background(m_cell1, "F8FAFC")
    set_cell_margins(m_cell1, top=60, bottom=60, left=100, right=100)
    p_meta1 = m_cell1.paragraphs[0]
    p_meta1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_meta1.paragraph_format.space_before = Pt(0)
    p_meta1.paragraph_format.space_after = Pt(0)
    r_m1 = p_meta1.add_run("CLASSIFICATION: STRICTLY CONFIDENTIAL")
    r_m1.font.size = Pt(8.5)
    r_m1.bold = True
    r_m1.font.color.rgb = RGBColor(71, 85, 105)
    
    doc.add_paragraph() # Spacer

    # 4. Parse Markdown Content
    lines = markdown_text.split('\n')
    i = 0
    total_lines = len(lines)
    
    while i < total_lines:
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Skip redundant main title lines (already in banner)
        if line.startswith('# ') or "CONFIDENTIAL" in line.upper():
            i += 1
            continue
            
        # --- Check for Markdown Table Block ---
        if line.startswith('|') and '|' in line[1:]:
            table_lines = []
            while i < total_lines and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
                
            # Filter separator rows
            parsed_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if cells and not all(re.match(r'^:?-+:?$', c) for c in cells):
                    parsed_rows.append(cells)
                    
            if parsed_rows:
                num_cols = max(len(r) for r in parsed_rows)
                for r in parsed_rows:
                    while len(r) < num_cols:
                        r.append("")
                        
                # Determine smart column widths
                col_widths = []
                if num_cols == 2:
                    col_widths = [Inches(2.5), Inches(4.6)]
                elif num_cols == 3:
                    col_widths = [Inches(2.4), Inches(2.0), Inches(2.7)]
                elif num_cols == 4:
                    col_widths = [Inches(2.8), Inches(1.1), Inches(1.0), Inches(2.2)]
                else:
                    w = 7.1 / num_cols
                    col_widths = [Inches(w)] * num_cols
                    
                word_table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
                word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                word_table.autofit = False
                set_table_borders(word_table, top="003366", bottom="003366", insideH="E2E8F0")
                
                # Apply column widths to table columns
                for c_idx, col in enumerate(word_table.columns):
                    col.width = col_widths[c_idx]
                
                # Format Header Row
                is_key_value_table = (num_cols == 2 and parsed_rows[0][0].lower() in ['metric', 'parameter', 'field'])
                
                header_row = word_table.rows[0]
                make_row_header(header_row)
                make_row_cant_split(header_row)
                
                for col_idx, text_val in enumerate(parsed_rows[0]):
                    cell = header_row.cells[col_idx]
                    cell.width = col_widths[col_idx]
                    set_cell_background(cell, "0A2540") # Dark Executive Navy
                    set_cell_margins(cell, top=130, bottom=130, left=140, right=140)
                    
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    
                    clean_h = sanitize_text(text_val.replace('**', ''))
                    if col_idx in [1, 2] and num_cols == 4:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    run = p.add_run(clean_h)
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    
                # Format Data Rows
                for row_idx in range(1, len(parsed_rows)):
                    row_obj = word_table.rows[row_idx]
                    make_row_cant_split(row_obj)
                    
                    row_cells_text = parsed_rows[row_idx]
                    is_total_row = any("TOTAL" in str(c).upper() for c in row_cells_text)
                    is_even = (row_idx % 2 == 0)
                    
                    for col_idx, text_val in enumerate(row_cells_text):
                        cell = row_obj.cells[col_idx]
                        cell.width = col_widths[col_idx]
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        
                        if is_total_row:
                            row_bg = "1E3A8A" # Highlight Blue for Total Score
                        elif num_cols == 2 and col_idx == 0:
                            row_bg = "F1F5F9" # Light slate for 2-column key column
                        else:
                            row_bg = "F8FAFC" if is_even else "FFFFFF"
                            
                        set_cell_background(cell, row_bg)
                        set_cell_margins(cell, top=90, bottom=90, left=140, right=140)
                        
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        
                        if is_total_row:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [1, 2] else WD_ALIGN_PARAGRAPH.LEFT
                            clean_t = sanitize_text(text_val.replace('**', ''))
                            run = p.add_run(clean_t)
                            run.bold = True
                            run.font.size = Pt(9.5)
                            run.font.color.rgb = RGBColor(255, 255, 255)
                        elif num_cols == 2 and col_idx == 0:
                            clean_label = sanitize_text(text_val.replace('**', ''))
                            run = p.add_run(clean_label)
                            run.bold = True
                            run.font.size = Pt(9.0)
                            run.font.color.rgb = RGBColor(15, 23, 42)
                        else:
                            if col_idx in [1, 2] and num_cols == 4:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                            add_formatted_runs(p, text_val, base_font_size=Pt(9.0))
                            
                doc.add_paragraph() # Spacer after table
            continue
            
        # --- Section Headings (##) ---
        if line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            clean_t = clean_section_title(line[3:].strip())
            run = p.add_run(clean_t)
            run.font.size = Pt(11.5)
            run.bold = True
            run.font.color.rgb = RGBColor(10, 37, 64) # Navy Heading
            
        # --- Sub-Headings (###) ---
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            
            clean_t = sanitize_text(line[4:].strip())
            run = p.add_run(clean_t)
            run.font.size = Pt(10.5)
            run.bold = True
            run.font.color.rgb = RGBColor(30, 64, 175) # Blue Subheading
            
        # --- Grouped Callout Box (> ...) ---
        elif line.startswith('> '):
            callout_lines = []
            while i < total_lines and lines[i].strip().startswith('> '):
                callout_lines.append(lines[i].strip()[2:].strip())
                i += 1
                
            callout_table = doc.add_table(rows=1, cols=1)
            callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            callout_table.autofit = False
            c_cell = callout_table.cell(0, 0)
            c_cell.width = Inches(7.1)
            set_cell_background(c_cell, "F1F5F9") # Soft Ice Blue
            set_cell_margins(c_cell, top=140, bottom=140, left=200, right=160)
            
            # Thick Navy Left Border
            tcPr = c_cell._element.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:left w:val="single" w:sz="24" w:space="0" w:color="003366"/>'
                f'<w:top w:val="none"/>'
                f'<w:bottom w:val="none"/>'
                f'<w:right w:val="none"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(borders)
            
            for c_idx, c_text in enumerate(callout_lines):
                if c_idx == 0:
                    cp = c_cell.paragraphs[0]
                else:
                    cp = c_cell.add_paragraph()
                cp.paragraph_format.space_before = Pt(2)
                cp.paragraph_format.space_after = Pt(2)
                add_formatted_runs(cp, c_text, base_font_size=Pt(9.5), default_color=RGBColor(15, 23, 42))
                
            doc.add_paragraph() # Spacer after callout box
            continue
            
        # --- Bullet List Items ---
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            content = line[2:].strip()
            add_formatted_runs(p, content, base_font_size=Pt(9.5))
            
        # --- Numbered List Items ---
        elif re.match(r'^\d+\.\s+', line):
            num_match = re.match(r'^(\d+\.)\s+(.*)', line)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            content = num_match.group(2)
            add_formatted_runs(p, content, base_font_size=Pt(9.5))
            
        # --- Regular Paragraph ---
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_runs(p, line, base_font_size=Pt(9.5))
            
        i += 1

    # 5. Executive Signature Block at the End
    doc.add_paragraph() # Spacer
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    make_row_cant_split(sig_table.rows[0])
    
    col_w = Inches(2.36)
    for cell in sig_table.rows[0].cells:
        cell.width = col_w
        set_cell_margins(cell, top=120, bottom=120, left=80, right=80)
        
    s0 = sig_table.cell(0, 0).paragraphs[0]
    s0.paragraph_format.space_before = Pt(0)
    s0.paragraph_format.space_after = Pt(0)
    s0.add_run("____________________________\n").font.color.rgb = RGBColor(148, 163, 184)
    s0.add_run("Prepared & Evaluated By:\n").bold = True
    s0.add_run("Credit Underwriter / AI Engine\nCentral Bank of India").font.size = Pt(8.5)
    
    s1 = sig_table.cell(0, 1).paragraphs[0]
    s1.paragraph_format.space_before = Pt(0)
    s1.paragraph_format.space_after = Pt(0)
    s1.add_run("____________________________\n").font.color.rgb = RGBColor(148, 163, 184)
    s1.add_run("Reviewed & Verified By:\n").bold = True
    s1.add_run("Branch Credit Manager\nCredit Operations").font.size = Pt(8.5)
    
    s2 = sig_table.cell(0, 2).paragraphs[0]
    s2.paragraph_format.space_before = Pt(0)
    s2.paragraph_format.space_after = Pt(0)
    s2.add_run("____________________________\n").font.color.rgb = RGBColor(148, 163, 184)
    s2.add_run("Sanctioned / Approved By:\n").bold = True
    s2.add_run("Zonal Credit Sanction Committee\nExecutive Authority").font.size = Pt(8.5)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
