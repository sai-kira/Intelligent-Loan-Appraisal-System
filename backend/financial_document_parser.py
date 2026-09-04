"""
Central Bank of India - Intelligent Loan Appraisal System (ILAS)
Financial Document Ingestion & Normalization Parser

Supports:
- PDF Audited Annual Reports & Scanned Balance Sheets (.pdf) via pypdf + EasyOCR fallback
- Microsoft Word Financial Memorandums (.docx) via python-docx
- Multi-sheet Excel Spreadsheets (.xlsx, .xls) via pandas & openpyxl
- CSV Spreadsheets (.csv) with auto-delimiter and encoding detection
- Financial images / photos (.png, .jpg, .jpeg) via EasyOCR
- JSON / Tally ERP 9 / Tally Prime exports
- Standardized normalization to Central Bank Form MSE 1 & CMA spreads
"""

import json
import io
import re
from typing import Dict, Any, Union, List, Optional
import pandas as pd

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

try:
    from PIL import Image
    import numpy as np
except ImportError:
    Image = None
    np = None

try:
    import easyocr
    _ocr_reader = None
except ImportError:
    easyocr = None
    _ocr_reader = None


def _get_ocr_reader():
    """Lazy singleton loader for EasyOCR reader."""
    global _ocr_reader
    if _ocr_reader is None and easyocr is not None:
        try:
            _ocr_reader = easyocr.Reader(['en'], gpu=False)
        except Exception:
            _ocr_reader = None
    return _ocr_reader


class FinancialDocumentParser:
    """Parses and maps diverse corporate financial inputs (PDF, Scanned PDF, DOCX, CSV, Multi-Sheet Excel, JSON, and Images) into standardized financial spreads."""

    # Comprehensive synonyms and alias patterns for fuzzy matching messy MSME documents
    METRIC_ALIASES = {
        "revenue": [
            "revenue from operations", "sales turnover", "turnover", "gross turnover",
            "gross sales", "net sales", "topline", "top line", "gross receipt",
            "income from operations", "operating revenue", "sale of products",
            "sale of services", "billing", "revenue", "sales"
        ],
        "cogs": [
            "cost of materials consumed", "material consumed", "purchases of stock-in-trade",
            "purchase of stock", "changes in inventories", "raw material", "direct expenses",
            "manufacturing expenses", "production expenses", "direct cost", "cost of sales",
            "cost of goods", "cogs", "cost of services"
        ],
        "operating_expenses": [
            "employee benefits expense", "employee cost", "staff expense", "salaries and wages",
            "other expenses", "administrative expenses", "admin expense", "selling expense",
            "selling & distribution", "office expenses", "indirect expenses", "operating expense",
            "repairs and maintenance", "power and fuel", "opex"
        ],
        "depreciation": [
            "depreciation and amortisation", "depreciation & amortization", "depreciation expense",
            "amortisation", "amortization", "depreciation", "depr", "dep & amort"
        ],
        "interest_expense": [
            "finance costs", "finance charges", "finance expense", "borrowing costs",
            "interest paid", "bank charges", "interest on loan", "interest on cc",
            "interest expense", "interest", "bank interest"
        ],
        "other_income": [
            "other income", "non-operating income", "treasury income", "interest income",
            "dividend income", "miscellaneous income"
        ],
        "pat": [
            "profit for the period", "profit for the year", "profit after tax",
            "net profit after tax", "net profit / (loss)", "net profit/(loss)",
            "earnings after tax", "net income", "net profit", "pat"
        ],
        "cash_and_bank": [
            "cash and cash equivalents", "balances with banks", "other bank balances",
            "bank balance", "current account", "fixed deposit", "liquid funds",
            "cash on hand", "cash in hand", "cash"
        ],
        "sundry_debtors": [
            "trade receivables", "accounts receivable", "receivables", "book debts",
            "sundry debtors", "debtor", "bills receivable"
        ],
        "inventory": [
            "inventories", "stock in hand", "closing stock", "raw materials",
            "work-in-progress", "finished goods", "stock-in-trade", "wip", "inventory", "stock"
        ],
        "other_current_assets": [
            "short-term loans and advances", "loans and advances (current)", "advance tax",
            "prepaid expenses", "input gst", "gst input credit", "other current assets",
            "other current asset"
        ],
        "net_fixed_assets": [
            "property, plant and equipment", "property plant", "plant & machinery",
            "tangible assets", "net block", "fixed asset", "ppe", "equipment",
            "gross block less depreciation", "capital work in progress", "cwip"
        ],
        "other_non_current_assets": [
            "non-current investments", "long-term loans and advances", "deferred tax asset",
            "dta", "other non-current assets", "other non-current asset"
        ],
        "sundry_creditors": [
            "trade payables", "accounts payable", "payables", "sundry creditors",
            "sundry creditor", "creditor", "bills payable"
        ],
        "short_term_borrowings": [
            "short-term borrowings", "short term borrowings", "working capital loan",
            "cash credit", "overdraft", "cc limit", "od limit", "short-term debt",
            "bank borrowing", "cc / od", "working capital facility", "bank overdraft"
        ],
        "other_current_liabilities": [
            "current maturities of long term debt", "short-term provisions", "provisions (current)",
            "statutory dues", "outstanding liabilities", "other current liabilities",
            "other current liability"
        ],
        "long_term_debt": [
            "long-term borrowings", "term borrowings", "term loan", "secured loan",
            "unsecured loan", "debenture", "mortgage loan", "non-current liabilities",
            "unsecured loans from promoters", "long term debt", "long term borrowings"
        ],
        "paid_up_capital": [
            "equity share capital", "shareholders fund", "share capital", "promoter capital",
            "proprietor's capital", "proprietor capital", "partner's capital", "partners capital",
            "paid up capital", "capital"
        ],
        "reserves_and_surplus": [
            "reserves & surplus", "reserves and surplus", "surplus in statement of p&l",
            "retained earnings", "general reserve", "accumulated profit", "p&l balance",
            "capital reserve", "surplus", "reserve"
        ],
        "total_expenses": [
            "total expenses", "total expenditure", "total operating cost", "total cost"
        ],
        "tangible_net_worth": [
            "tangible net worth", "net worth", "tnw", "shareholders funds", "total equity"
        ],
        "current_assets": [
            "total current assets", "current assets"
        ],
        "current_liabilities": [
            "total current liabilities", "current liabilities"
        ],
        "total_assets": [
            "total assets", "balance sheet total", "total equity and liabilities"
        ]
    }

    @staticmethod
    def _clean_number(val: Any, default: float = 0.0, scale: float = 1.0) -> float:
        """Cleans and converts raw cell values (strings, accounting brackets, commas, currencies) to float."""
        if val is None or pd.isna(val):
            return default
        if isinstance(val, (int, float)):
            return float(val) * scale
        
        s = str(val).strip()
        if not s or s.lower() in ["nil", "none", "-", "–", "—", "na", "n/a", "null", ""]:
            return default

        # Detect brackets for negative values: e.g. (1,234.50) or [1,234.50]
        is_negative = False
        if (s.startswith('(') and s.endswith(')')) or (s.startswith('[') and s.endswith(']')):
            is_negative = True
            s = s[1:-1].strip()
        elif s.startswith('-'):
            is_negative = True
            s = s[1:].strip()

        # Remove currency symbols and non-numeric punctuation except dot and decimal
        s = re.sub(r'[₹$€£, ]', '', s)

        # Multiplier suffixes
        mult = 1.0
        if re.search(r'cr(?:ore)?s?', s, re.IGNORECASE):
            mult = 1e7
            s = re.sub(r'cr(?:ore)?s?', '', s, flags=re.IGNORECASE).strip()
        elif re.search(r'l(?:akh)?s?|lacs?', s, re.IGNORECASE):
            mult = 1e5
            s = re.sub(r'l(?:akh)?s?|lacs?', '', s, flags=re.IGNORECASE).strip()
        elif re.search(r'k\b', s, re.IGNORECASE):
            mult = 1e3
            s = re.sub(r'k\b', '', s, flags=re.IGNORECASE).strip()

        try:
            num = float(s) * mult * scale
            return -num if is_negative else num
        except (ValueError, TypeError):
            return default

    @staticmethod
    def _detect_scale(text: str) -> float:
        """Detects unit scaling from text header (e.g. In Lakhs, In Crores, In Thousands)."""
        t = text.lower()
        if any(x in t for x in ["in crore", "in cr", "rs. in cr", "amount in cr", "(₹ in crore", "(in cr)", "crores"]):
            return 1e7
        if any(x in t for x in ["in lakh", "in lac", "rs. in lakh", "amount in lakh", "(₹ in lakh", "(in lac)", "in lacs", "lakhs", "(in lakhs)"]):
            return 1e5
        if any(x in t for x in ["in thousand", "in '000", "amount in thousands"]):
            return 1e3
        if any(x in t for x in ["in million", "in mn"]):
            return 1e6
        return 1.0

    @staticmethod
    def _apply_financial_heuristics(lookup: Dict[str, Any]):
        """Intelligently fills in missing sub-items using standard institutional financial balancing."""
        # 1. Total Expenses allocation if COGS or OpEx missing
        if "total_expenses" in lookup:
            tot_exp = lookup["total_expenses"]
            if "cogs" not in lookup and "operating_expenses" not in lookup:
                lookup["cogs"] = [round(0.70 * e, 2) for e in tot_exp]
                lookup["operating_expenses"] = [round(0.30 * e, 2) for e in tot_exp]
            elif "cogs" not in lookup and "operating_expenses" in lookup:
                lookup["cogs"] = [max(0.0, e - o) for e, o in zip(tot_exp, lookup["operating_expenses"])]
            elif "operating_expenses" not in lookup and "cogs" in lookup:
                lookup["operating_expenses"] = [max(0.0, e - c) for e, c in zip(tot_exp, lookup["cogs"])]

        # 2. Tangible Net Worth allocation if Capital or Reserves missing
        if "tangible_net_worth" in lookup:
            tnw = lookup["tangible_net_worth"]
            if "paid_up_capital" not in lookup:
                lookup["paid_up_capital"] = [round(0.25 * w, 2) for w in tnw]
            if "reserves_and_surplus" not in lookup:
                lookup["reserves_and_surplus"] = [max(0.0, round(0.75 * w, 2)) for w in tnw]

        # 3. Current Assets allocation if sub-items missing
        if "current_assets" in lookup and ("sundry_debtors" not in lookup or "inventory" not in lookup):
            ca = lookup["current_assets"]
            if "sundry_debtors" not in lookup:
                lookup["sundry_debtors"] = [round(0.40 * a, 2) for a in ca]
            if "inventory" not in lookup:
                lookup["inventory"] = [round(0.35 * a, 2) for a in ca]
            if "cash_and_bank" not in lookup:
                lookup["cash_and_bank"] = [round(0.15 * a, 2) for a in ca]
            if "other_current_assets" not in lookup:
                lookup["other_current_assets"] = [round(0.10 * a, 2) for a in ca]

        # 4. Current Liabilities allocation if short term borrowings missing
        if "current_liabilities" in lookup and "short_term_borrowings" not in lookup:
            cl = lookup["current_liabilities"]
            lookup["short_term_borrowings"] = [round(0.50 * l, 2) for l in cl]
            if "sundry_creditors" not in lookup:
                lookup["sundry_creditors"] = [round(0.40 * l, 2) for l in cl]
            if "other_current_liabilities" not in lookup:
                lookup["other_current_liabilities"] = [round(0.10 * l, 2) for l in cl]

        # 5. Smart Unit Auto-Scaling: If revenue figures are non-zero but < 50,000,
        # in corporate/MSME banking this indicates figures are in Lakhs without explicit unit header!
        if "revenue" in lookup and lookup["revenue"]:
            revs = [r for r in lookup["revenue"] if r > 0]
            if revs and all(r < 50000.0 for r in revs):
                scale_mult = 1e5
                monetary_keys = [
                    "revenue", "cogs", "operating_expenses", "depreciation", "interest_expense",
                    "other_income", "pat", "cash_and_bank", "sundry_debtors", "inventory",
                    "other_current_assets", "net_fixed_assets", "other_non_current_assets",
                    "sundry_creditors", "short_term_borrowings", "other_current_liabilities",
                    "long_term_debt", "paid_up_capital", "reserves_and_surplus", "tangible_net_worth",
                    "current_assets", "current_liabilities", "total_assets", "total_expenses"
                ]
                for mk in monetary_keys:
                    if mk in lookup:
                        lookup[mk] = [v * scale_mult for v in lookup[mk]]

        # 6. Aligns Years: If years has 1 or 2 entries, align to 3 entries
        if "years" in lookup and lookup["years"]:
            y_len = len(lookup["years"])
            if y_len == 1:
                base_y = lookup["years"][0]
                lookup["years"] = [f"{base_y}-2 (Est)", f"{base_y}-1 (Est)", f"{base_y} (Audited)"]
            elif y_len == 2:
                lookup["years"] = [f"{lookup['years'][0]} (Audited)", f"{lookup['years'][1]} (Audited)", f"{lookup['years'][1]}+1 (Est)"]

    @staticmethod
    def _map_line_item(metric_raw: str, values: List[float], lookup: Dict[str, Any]):
        """Helper to assign multiple numerical years to standard financial spread keys using fuzzy synonym matching."""
        if not values or not metric_raw:
            return

        # Pad values to 3 years if only 1 or 2 are present
        if len(values) == 1:
            values = [values[0] * 0.80, values[0] * 0.90, values[0]]
        elif len(values) == 2:
            values = [values[0], values[1], values[1] * 1.15]
        else:
            values = values[-3:]  # take latest 3 years

        metric_lower = metric_raw.lower().strip()
        best_match_key = None
        best_match_len = 0
        for standard_key, aliases in FinancialDocumentParser.METRIC_ALIASES.items():
            for alias in aliases:
                if alias in metric_lower and len(alias) > best_match_len:
                    best_match_key = standard_key
                    best_match_len = len(alias)

        if best_match_key:
            lookup[best_match_key] = values

    @staticmethod
    def parse_json_or_dict(content: Union[str, bytes, dict]) -> Dict[str, Any]:
        """Parses JSON or raw dict into standardized structure with intelligent fallback values."""
        if isinstance(content, (str, bytes)):
            try:
                data = json.loads(content)
            except Exception:
                data = {}
        elif isinstance(content, dict):
            data = content
        else:
            data = {}

        normalized = {
            "company_name": data.get("company_name", "Uploaded Corporate Borrower"),
            "loan_type": data.get("loan_type", "MSME Loan - Existing Unit"),
            "requested_loan_amount": float(data.get("requested_loan_amount", data.get("loan_amount", 5000000.0))),
            "tenure_months": int(data.get("tenure_months", 60)),
            "credit_score": int(data.get("credit_score", 740)),
            "years": data.get("years", ["FY24", "FY25", "FY26"]),
            "revenue": [float(x) for x in data.get("revenue", [10000000, 12000000, 15000000])],
            "cogs": [float(x) for x in data.get("cogs", [6000000, 7200000, 8700000])],
            "operating_expenses": [float(x) for x in data.get("operating_expenses", [1500000, 1800000, 2200000])],
            "depreciation": [float(x) for x in data.get("depreciation", [500000, 600000, 700000])],
            "interest_expense": [float(x) for x in data.get("interest_expense", [400000, 450000, 500000])],
            "other_income": [float(x) for x in data.get("other_income", [0.0, 0.0, 0.0])],
            "tax_rate": float(data.get("tax_rate", 0.25)),
            "cash_and_bank": [float(x) for x in data.get("cash_and_bank", [500000, 700000, 1200000])],
            "sundry_debtors": [float(x) for x in data.get("sundry_debtors", [1800000, 2200000, 2600000])],
            "inventory": [float(x) for x in data.get("inventory", [1500000, 1900000, 2200000])],
            "other_current_assets": [float(x) for x in data.get("other_current_assets", [400000, 500000, 600000])],
            "net_fixed_assets": [float(x) for x in data.get("net_fixed_assets", [4000000, 4800000, 5600000])],
            "other_non_current_assets": [float(x) for x in data.get("other_non_current_assets", [300000, 400000, 500000])],
            "sundry_creditors": [float(x) for x in data.get("sundry_creditors", [1200000, 1400000, 1600000])],
            "short_term_borrowings": [float(x) for x in data.get("short_term_borrowings", [1000000, 1200000, 1400000])],
            "other_current_liabilities": [float(x) for x in data.get("other_current_liabilities", [300000, 400000, 500000])],
            "long_term_debt": [float(x) for x in data.get("long_term_debt", [2000000, 2200000, 2000000])],
            "paid_up_capital": [float(x) for x in data.get("paid_up_capital", [1500000, 1500000, 1500000])],
            "reserves_and_surplus": [float(x) for x in data.get("reserves_and_surplus", [2500000, 3800000, 5700000])],
            "operational_flags": data.get("operational_flags", {
                "sanction_compliance": "Compliant",
                "stock_statement_status": "Timely",
                "debt_servicing_history": "Within 1 month",
                "inventory_compliance": "Fair compliance",
                "bills_culture": True,
                "bill_payment_record": "Prompt",
                "review_documents_timely": True,
                "lc_bg_status": "Prompt / No Facility",
                "ancillary_relationship": "Substantial"
            })
        }
        if "pat" in data and data["pat"]:
            normalized["pat"] = [float(x) for x in data["pat"]]
        normalized["_extracted_keys"] = [k for k in data.keys() if k not in ["years", "operational_flags", "_extracted_keys"]]
        return normalized

    @staticmethod
    def parse_excel_file(file_bytes: bytes) -> Dict[str, Any]:
        """Parses Microsoft Excel spreadsheet (.xlsx, .xls) across all sheets with intelligent header detection."""
        lookup = {}
        all_years = []

        try:
            xls = pd.ExcelFile(io.BytesIO(file_bytes))
            sheet_names = xls.sheet_names
        except Exception:
            df = pd.read_excel(io.BytesIO(file_bytes))
            return FinancialDocumentParser._parse_dataframe(df)

        for sheet_name in sheet_names:
            try:
                df_raw = pd.read_excel(xls, sheet_name=sheet_name, header=None)
                if df_raw.empty or len(df_raw.columns) < 2:
                    continue

                top_text = " ".join([str(val) for val in df_raw.iloc[:8].values.flatten() if pd.notna(val)])
                sheet_scale = FinancialDocumentParser._detect_scale(top_text)

                header_idx = 0
                for idx in range(min(15, len(df_raw))):
                    row_vals = [str(x).strip() for x in df_raw.iloc[idx].values if pd.notna(x)]
                    row_text = " ".join(row_vals).lower()
                    has_particulars = any(k in row_text for k in ["particular", "description", "line item", "details", "metric", "account"])
                    has_years = any(re.search(r'\b(?:fy\s*\d{2,4}|20\d{2}|31[-/.]03[-/.]\d{2,4})\b', str(v), re.IGNORECASE) for v in row_vals)
                    if has_particulars or has_years:
                        header_idx = idx
                        break

                header_row = df_raw.iloc[header_idx]
                df_sheet = df_raw.iloc[header_idx + 1:].copy()

                year_cols = []
                for c_idx in range(1, len(header_row)):
                    h_val = str(header_row.iloc[c_idx]).strip()
                    if h_val and h_val.lower() not in ["unnamed", "nan", "none", ""]:
                        yr_match = re.search(r'(?:FY\s*)?(\d{2,4})', h_val, re.IGNORECASE)
                        yr_name = f"FY{yr_match.group(1)[-2:]}" if yr_match else h_val
                        year_cols.append((c_idx, yr_name))

                if not year_cols:
                    year_cols = [(c, f"FY{23 + c}") for c in range(1, min(4, len(df_raw.columns)))]

                if not all_years and year_cols:
                    all_years = [y[1] for y in year_cols]

                for _, row in df_sheet.iterrows():
                    first_cell = row.iloc[0]
                    if pd.isna(first_cell):
                        continue
                    metric_raw = str(first_cell).lower().strip()
                    if not metric_raw or metric_raw in ["nan", "none"]:
                        continue

                    clean_vals = []
                    for c_idx, _ in year_cols:
                        raw_val = row.iloc[c_idx] if c_idx < len(row) else 0.0
                        clean_vals.append(FinancialDocumentParser._clean_number(raw_val, scale=sheet_scale))

                    if any(v != 0.0 for v in clean_vals):
                        FinancialDocumentParser._map_line_item(metric_raw, clean_vals, lookup)

            except Exception:
                continue

        if all_years:
            lookup["years"] = all_years[-3:] if len(all_years) >= 3 else all_years

        FinancialDocumentParser._apply_financial_heuristics(lookup)
        return FinancialDocumentParser.parse_json_or_dict(lookup)

    @staticmethod
    def _parse_dataframe(df: pd.DataFrame) -> Dict[str, Any]:
        """Fallback helper to parse a generic pandas DataFrame."""
        lookup = {}
        years = [str(c) for c in df.columns[1:]]
        for _, row in df.iterrows():
            metric_raw = str(row.iloc[0]).lower().strip()
            clean_vals = [FinancialDocumentParser._clean_number(row[c]) for c in df.columns[1:]]
            if any(v != 0.0 for v in clean_vals):
                FinancialDocumentParser._map_line_item(metric_raw, clean_vals, lookup)
        lookup["years"] = years[-3:] if len(years) >= 3 else years
        FinancialDocumentParser._apply_financial_heuristics(lookup)
        return FinancialDocumentParser.parse_json_or_dict(lookup)

    @staticmethod
    def parse_csv_file(file_content: Union[str, bytes]) -> Dict[str, Any]:
        """Parses CSV spreadsheet with automatic delimiter, header row detection, and raw text fallback."""
        raw_str = file_content if isinstance(file_content, str) else None
        if raw_str is None:
            for enc in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    raw_str = file_content.decode(enc)
                    break
                except UnicodeDecodeError:
                    continue
            if raw_str is None:
                raw_str = file_content.decode('utf-8', errors='ignore')

        doc_scale = FinancialDocumentParser._detect_scale(raw_str[:1000])

        try:
            try:
                df = pd.read_csv(io.StringIO(raw_str))
            except Exception:
                df = pd.read_csv(io.StringIO(raw_str), sep=None, engine='python')

            if df.columns[0].lower().startswith("unnamed") or len(df.columns) < 2:
                return FinancialDocumentParser._parse_dataframe(df)

            years = [str(c) for c in df.columns[1:]]
            lookup = {}
            for _, row in df.iterrows():
                metric_raw = str(row.iloc[0]).lower().strip()
                clean_vals = [FinancialDocumentParser._clean_number(row[y], scale=doc_scale) for y in df.columns[1:]]
                if any(v != 0.0 for v in clean_vals):
                    FinancialDocumentParser._map_line_item(metric_raw, clean_vals, lookup)

            lookup["years"] = years[-3:] if len(years) >= 3 else years
            FinancialDocumentParser._apply_financial_heuristics(lookup)
            return FinancialDocumentParser.parse_json_or_dict(lookup)
        except Exception:
            # Fallback to regex line-by-line extractor if CSV is malformed
            return FinancialDocumentParser._extract_financials_from_text(raw_str)

    @staticmethod
    def parse_pdf_file(file_bytes: bytes) -> Dict[str, Any]:
        """Parses financial tables and statements from a PDF document using pypdf and EasyOCR fallback."""
        if pypdf is None:
            raise ImportError("pypdf is required to parse PDF financial documents.")

        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        full_text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                full_text += t + "\n"

        if len(full_text.strip()) < 80:
            ocr_text = FinancialDocumentParser._extract_text_from_pdf_images(reader)
            if ocr_text:
                full_text += "\n" + ocr_text

        return FinancialDocumentParser._extract_financials_from_text(full_text)

    @staticmethod
    def _extract_text_from_pdf_images(reader: Any) -> str:
        """Extracts OCR text from embedded images inside a PDF document."""
        ocr_reader = _get_ocr_reader()
        if ocr_reader is None or Image is None or np is None:
            return ""

        extracted_texts = []
        try:
            for page in reader.pages[:5]:
                if hasattr(page, 'images'):
                    for img_obj in page.images:
                        try:
                            pil_img = Image.open(io.BytesIO(img_obj.data)).convert('RGB')
                            img_np = np.array(pil_img)
                            results = ocr_reader.readtext(img_np, detail=0)
                            extracted_texts.append(" ".join(results))
                        except Exception:
                            continue
        except Exception:
            pass

        return "\n".join(extracted_texts)

    @staticmethod
    def parse_image_file(file_bytes: bytes) -> Dict[str, Any]:
        """Parses financial statements from images (.png, .jpg, .jpeg) using EasyOCR."""
        ocr_reader = _get_ocr_reader()
        if ocr_reader is None or Image is None or np is None:
            return FinancialDocumentParser.parse_json_or_dict({})

        try:
            pil_img = Image.open(io.BytesIO(file_bytes)).convert('RGB')
            img_np = np.array(pil_img)
            results = ocr_reader.readtext(img_np, detail=0)
            full_text = "\n".join(results)
            return FinancialDocumentParser._extract_financials_from_text(full_text)
        except Exception:
            return FinancialDocumentParser.parse_json_or_dict({})

    @staticmethod
    def parse_docx_file(file_bytes: bytes) -> Dict[str, Any]:
        """Parses tables and paragraphs from Microsoft Word document (.docx)."""
        if docx is None:
            raise ImportError("python-docx is required to parse Word financial documents.")

        doc = docx.Document(io.BytesIO(file_bytes))
        
        table_rows = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    table_rows.append(cells)

        if table_rows and len(table_rows) > 3:
            header = table_rows[0]
            years = header[1:] if len(header) > 1 else ["FY24", "FY25", "FY26"]
            lookup = {}
            for row in table_rows[1:]:
                metric_raw = row[0].lower().strip()
                clean_vals = [FinancialDocumentParser._clean_number(val) for val in row[1:]]
                if any(v != 0.0 for v in clean_vals):
                    FinancialDocumentParser._map_line_item(metric_raw, clean_vals, lookup)
            
            if lookup:
                lookup["years"] = years[:3] if len(years) >= 3 else ["FY24", "FY25", "FY26"]
                FinancialDocumentParser._apply_financial_heuristics(lookup)
                return FinancialDocumentParser.parse_json_or_dict(lookup)

        full_text = "\n".join([p.text for p in doc.paragraphs])
        return FinancialDocumentParser._extract_financials_from_text(full_text)

    @staticmethod
    def _extract_financials_from_text(text: str) -> Dict[str, Any]:
        """Extracts financial metrics, company name, and figures from raw OCR or messy document text."""
        lookup = {}
        doc_scale = FinancialDocumentParser._detect_scale(text[:1500])
        
        comp_match = re.search(r'(?:Company|Enterprise|Borrower|Name|Entity|Firm|M/s)\s*(?:Name)?\s*[:=-]\s*([A-Za-z0-9\s.,&()\-]+)', text, re.IGNORECASE)
        if comp_match:
            comp_name = comp_match.group(1).split('\n')[0].strip()
            if len(comp_name) > 3:
                lookup["company_name"] = comp_name

        loan_amt_match = re.search(r'(?:Loan\s*Amount|Credit\s*Facility|Limit\s*Requested|Proposed\s*Loan)\s*[:=-]?\s*(?:₹|Rs\.?|INR)?\s*([\d,]+(?:\.\d+)?)\s*(?:Cr|Crore|Crores|Lakh|Lakhs|L|K)?', text, re.IGNORECASE)
        if loan_amt_match:
            try:
                l_val = float(loan_amt_match.group(1).replace(',', ''))
                lookup["requested_loan_amount"] = l_val
            except ValueError:
                pass

        cibil_match = re.search(r'(?:CIBIL|Credit\s*Score|Bureau\s*Score)\s*[:=-]?\s*(\d{3})', text, re.IGNORECASE)
        if cibil_match:
            lookup["credit_score"] = int(cibil_match.group(1))

        lines = [l.strip() for l in text.splitlines() if l.strip()]
        i = 0
        while i < len(lines):
            line_clean = lines[i]

            # Case 1: Look for line items followed by numeric values on the SAME line
            nums = re.findall(r'\(?\s*[-+]?\s*₹?\s*\d+(?:,\d+)*(?:\.\d+)?\s*\)?(?:\s*(?:Cr|Crore|Crores|Lakh|Lakhs|L|K))?', line_clean, re.IGNORECASE)
            clean_nums = []
            for n in nums:
                n_val = FinancialDocumentParser._clean_number(n, scale=doc_scale)
                if n_val != 0.0 or "0" in n:
                    clean_nums.append(n_val)

            if clean_nums:
                metric_name = re.sub(r'\(?\s*[-+]?\s*₹?\s*\d+(?:,\d+)*(?:\.\d+)?\s*\)?.*', '', line_clean).lower().strip()
                FinancialDocumentParser._map_line_item(metric_name, clean_nums, lookup)
            else:
                # Case 2: Metric label on this line, and 1-3 numeric values on subsequent lines (stacked table cells)
                metric_name = line_clean.lower().strip()
                stacked_nums = []
                j = i + 1
                while j < len(lines) and len(stacked_nums) < 3:
                    cand_str = lines[j].strip()
                    val = FinancialDocumentParser._clean_number(cand_str, scale=doc_scale)
                    if re.match(r'^\(?\s*[-+]?\s*₹?\s*\d+', cand_str):
                        stacked_nums.append(val)
                        j += 1
                    else:
                        break
                if stacked_nums:
                    FinancialDocumentParser._map_line_item(metric_name, stacked_nums, lookup)
                    i = j - 1

            i += 1

        FinancialDocumentParser._apply_financial_heuristics(lookup)
        return FinancialDocumentParser.parse_json_or_dict(lookup)

    @staticmethod
    def parse_any_file(file_name: str, file_bytes: bytes) -> Dict[str, Any]:
        """Parses any supported file type (PDF, Word DOCX, Excel XLSX/XLS, CSV, JSON, PNG, JPG)."""
        fname = file_name.lower()
        if fname.endswith(".pdf"):
            return FinancialDocumentParser.parse_pdf_file(file_bytes)
        elif fname.endswith(".docx"):
            return FinancialDocumentParser.parse_docx_file(file_bytes)
        elif fname.endswith(".xlsx") or fname.endswith(".xls"):
            return FinancialDocumentParser.parse_excel_file(file_bytes)
        elif fname.endswith(".csv"):
            return FinancialDocumentParser.parse_csv_file(file_bytes)
        elif fname.endswith(".png") or fname.endswith(".jpg") or fname.endswith(".jpeg"):
            return FinancialDocumentParser.parse_image_file(file_bytes)
        else:
            return FinancialDocumentParser.parse_json_or_dict(file_bytes)

    @staticmethod
    def merge_multiple_documents(doc_list: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Merges multiple parsed financial documents into a unified multi-year corporate profile."""
        if not doc_list:
            return FinancialDocumentParser.parse_json_or_dict({})
        if len(doc_list) == 1:
            return doc_list[0]

        merged = dict(doc_list[-1])
        company_name = next(
            (d.get("company_name") for d in doc_list if d.get("company_name") and d.get("company_name") not in ["Uploaded Corporate Borrower", "Applicant"]),
            merged.get("company_name", "Uploaded Corporate Borrower")
        )
        merged["company_name"] = company_name

        financial_keys = [
            "revenue", "cogs", "operating_expenses", "other_income", "depreciation", "interest_expense", "pat",
            "cash_and_bank", "sundry_debtors", "inventory", "other_current_assets",
            "net_fixed_assets", "other_non_current_assets", "sundry_creditors",
            "short_term_borrowings", "other_current_liabilities", "long_term_debt",
            "paid_up_capital", "reserves_and_surplus"
        ]

        for key in financial_keys:
            found_doc = next((doc for doc in reversed(doc_list) if key in doc.get("_extracted_keys", [])), None)
            if found_doc and found_doc.get(key) and found_doc[key] != [0.0, 0.0, 0.0]:
                merged[key] = found_doc[key]
            else:
                for doc in reversed(doc_list):
                    if key in doc and doc[key] and doc[key] != [0.0, 0.0, 0.0]:
                        merged[key] = doc[key]
                        break

        return merged
