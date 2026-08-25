import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def sanitize_text(text):
    if not text:
        return ""
    # Remove XML-incompatible control characters
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout_box(doc, text_lines, bg_color="F0F4F8", border_color="002B49"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    for i, line in enumerate(text_lines):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
        run = p.add_run(sanitize_text(line))
        run.font.name = 'Consolas' if line.startswith("+") or line.startswith("|") else 'Calibri'
        run.font.size = Pt(8.5 if line.startswith("+") or line.startswith("|") else 10)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

def compile_academic_report():
    report_dir = os.path.dirname(os.path.abspath(__file__))
    output_docx = os.path.join(report_dir, "Central_Bank_of_India_ILAS_Academic_Report.docx")
    
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Header / Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Central Bank of India | Intelligent Loan Appraisal System (ILAS)")
        hrun.font.name = "Calibri"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Confidential — Academic & Institutional Credit Underwriting Dissertation")
        frun.font.name = "Calibri"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # Gather all markdown files in order
    files = sorted([f for f in os.listdir(report_dir) if f.endswith(".md")])
    if not files:
        print("No markdown report files found.")
        return

    first_file = True
    for fname in files:
        fpath = os.path.join(report_dir, fname)
        with open(fpath, "r", encoding="utf-8") as f:
            lines = f.readlines()

        if not first_file:
            doc.add_page_break()
        first_file = False

        in_code_block = False
        code_block_lines = []
        table_lines = []

        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Page Break
            if "<!-- PAGE BREAK -->" in line:
                doc.add_page_break()
                i += 1
                continue

            # Code block
            if line.startswith("```"):
                if in_code_block:
                    in_code_block = False
                    add_callout_box(doc, code_block_lines, bg_color="F8FAFC", border_color="002B49")
                    code_block_lines = []
                else:
                    in_code_block = True
                    code_block_lines = []
                i += 1
                continue

            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue

            # Table detection
            if line.startswith("|") and line.endswith("|"):
                table_lines.append(line)
                if i + 1 < len(lines) and lines[i+1].rstrip().startswith("|"):
                    i += 1
                    continue
                else:
                    render_markdown_table(doc, table_lines)
                    table_lines = []
                    i += 1
                    continue

            # Headings
            if line.startswith("# "):
                text = line[2:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(sanitize_text(text))
                run.font.name = "Arial"
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x2B, 0x49) # Navy
                i += 1
                continue

            if line.startswith("## "):
                text = line[3:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(sanitize_text(text))
                run.font.name = "Arial"
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x9E, 0x1B, 0x32) # Burgundy
                i += 1
                continue

            if line.startswith("### "):
                text = line[4:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(sanitize_text(text))
                run.font.name = "Calibri"
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
                i += 1
                continue

            # Ignore raw HTML divs
            if line.startswith("<div") or line.startswith("</div>") or line.startswith("<br") or line.startswith("---"):
                i += 1
                continue

            # Empty lines
            if not line.strip():
                i += 1
                continue

            # Regular paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            # Bullet list item
            if line.startswith("* ") or line.startswith("- "):
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                line = "• " + line[2:]
            elif re.match(r'^\d+\.\s', line):
                p.paragraph_format.left_indent = Inches(0.25)

            parse_inline_formatting(p, line)
            i += 1

    doc.save(output_docx)
    print(f"Successfully generated Word document at: {output_docx}")

def parse_inline_formatting(paragraph, text):
    text = sanitize_text(text)
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for token in tokens:
        if not token:
            continue
        token_clean = sanitize_text(token)
        if token_clean.startswith("**") and token_clean.endswith("**"):
            run = paragraph.add_run(token_clean[2:-2])
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        elif token_clean.startswith("*") and token_clean.endswith("*"):
            run = paragraph.add_run(token_clean[1:-1])
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.italic = True
        elif token_clean.startswith("`") and token_clean.endswith("`"):
            run = paragraph.add_run(token_clean[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x83, 0x18, 0x43)
        else:
            run = paragraph.add_run(token_clean)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

def render_markdown_table(doc, table_lines):
    if len(table_lines) < 2:
        return
    
    parsed_rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if any(set(c).issubset({'-', ':', ' '}) for c in cells):
            continue
        parsed_rows.append(cells)
        
    if not parsed_rows:
        return
        
    num_cols = len(parsed_rows[0])
    num_rows = len(parsed_rows)
    
    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    col_width = Inches(6.5 / max(1, num_cols))
    
    for r_idx, row in enumerate(tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = col_width
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            
            raw_text = parsed_rows[r_idx][c_idx] if c_idx < len(parsed_rows[r_idx]) else ""
            raw_clean = sanitize_text(raw_text)
            
            if r_idx == 0:
                set_cell_background(cell, "002B49")
                run = p.add_run(raw_clean.replace("**", ""))
                run.font.name = "Arial"
                run.font.size = Pt(9.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F9FAFB")
                else:
                    set_cell_background(cell, "FFFFFF")
                parse_inline_formatting(p, raw_clean)
                
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)

if __name__ == "__main__":
    compile_academic_report()
